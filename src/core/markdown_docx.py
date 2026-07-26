"""
Meeting Scribe — Markdown to .docx converter
Turns the Markdown returned by an LLM into a formatted Word document.

Supports the subset LLMs actually produce:
    # / ## / ### headings, - and * bullets, 1. numbered lists,
    **bold** and *italic* inline, | tables |, --- rules, > quotes.
"""
from __future__ import annotations

import re
import logging
from typing import List

logger = logging.getLogger(__name__)


def markdown_to_docx(markdown: str, output_path: str,
                     title: str = "", subtitle: str = "") -> str:
    """
    Convert Markdown text to a .docx file.

    Args:
        markdown: The Markdown source.
        output_path: Where to write the .docx.
        title: Optional document title inserted at the top.
        subtitle: Optional subtitle line under the title.

    Returns:
        The output path.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if title:
        h = doc.add_heading(title, level=0)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    lines = markdown.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Blank
        if not line.strip():
            i += 1
            continue

        # Table (needs a separator row underneath)
        if line.lstrip().startswith("|") and i + 1 < len(lines) \
                and re.match(r"^\s*\|[\s:\-|]+\|\s*$", lines[i + 1]):
            i = _add_table(doc, lines, i)
            continue

        # Horizontal rule
        if re.match(r"^\s*(-{3,}|\*{3,}|_{3,})\s*$", line):
            doc.add_paragraph()
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 4)
            h = doc.add_heading("", level=level)
            _add_inline(h, _strip_md(m.group(2)))
            from docx.shared import RGBColor as _RGB
            for run in h.runs:
                run.font.color.rgb = _RGB(0x1F, 0x38, 0x64)
            i += 1
            continue

        # Block quote
        if line.lstrip().startswith(">"):
            p = doc.add_paragraph(style="Intense Quote" if _has_style(doc, "Intense Quote") else None)
            _add_inline(p, line.lstrip().lstrip(">").strip())
            i += 1
            continue

        # Bullet list
        m = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            style_name = "List Bullet" if indent < 2 else "List Bullet 2"
            p = doc.add_paragraph(style=style_name if _has_style(doc, style_name) else "List Bullet")
            _add_inline(p, m.group(2))
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\s*)\d+[.)]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number" if _has_style(doc, "List Number") else None)
            _add_inline(p, m.group(2))
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        _add_inline(p, line)
        i += 1

    doc.save(output_path)
    logger.info(f"Markdown converted to docx: {output_path}")
    return output_path


def _has_style(doc, name: str) -> bool:
    try:
        doc.styles[name]
        return True
    except KeyError:
        return False


def _strip_md(text: str) -> str:
    return text.strip()


def _add_inline(paragraph, text: str) -> None:
    """Add text to a paragraph, honoring **bold**, *italic* and `code`."""
    tokens = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            run = paragraph.add_run(tok[1:-1])
            run.italic = True
        elif tok.startswith("`") and tok.endswith("`") and len(tok) > 2:
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(tok)


def _add_table(doc, lines: List[str], start: int) -> int:
    """Parse a Markdown table starting at `start`; returns the next line index."""
    def split_row(row: str) -> List[str]:
        return [c.strip() for c in row.strip().strip("|").split("|")]

    header = split_row(lines[start])
    i = start + 2  # skip header + separator
    rows = []
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        rows.append(split_row(lines[i]))
        i += 1

    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Light Grid Accent 1" if _has_style(doc, "Light Grid Accent 1") else "Table Grid"

    hdr = table.rows[0].cells
    for c, text in enumerate(header[:ncols]):
        hdr[c].text = ""
        _add_inline(hdr[c].paragraphs[0], text)
        for run in hdr[c].paragraphs[0].runs:
            run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for c in range(ncols):
            val = row[c] if c < len(row) else ""
            cells[c].text = ""
            _add_inline(cells[c].paragraphs[0], val)

    doc.add_paragraph()
    return i
