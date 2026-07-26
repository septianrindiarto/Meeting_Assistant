"""
Meeting Scribe — Starter Template Generator
Creates the 6 built-in .docx templates with proper Jinja2 placeholders.
Run this script once to generate templates into the templates/ folder.
"""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def set_cell_shading(cell, color):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_mom_template(output_dir):
    """Minutes of Meeting — formal template."""
    doc = Document()

    # Title
    title = doc.add_heading('Minutes of Meeting', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meeting info
    doc.add_paragraph('{{ meeting.title }}', style='Title')
    doc.add_paragraph('Date: {{ meeting.date }}  |  Duration: {{ meeting.duration }}')
    doc.add_paragraph('')

    # Attendees
    doc.add_heading('Attendees', level=2)
    doc.add_paragraph('{% for person in attendees %}• {{ person }}\n{% endfor %}')

    # Summary
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph('{{ summary }}')

    # Decisions
    doc.add_heading('Key Decisions', level=2)
    doc.add_paragraph('{% for decision in decisions %}{{ loop.index }}. {{ decision.description }} — {{ decision.speaker }}\n{% endfor %}')

    # Action Items
    doc.add_heading('Action Items', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Owner'
    hdr[1].text = 'Action'
    hdr[2].text = 'Due Date'

    doc.add_paragraph('{% for item in action_items %}')
    doc.add_paragraph('{{ item.owner }} | {{ item.description }} | {{ item.due_date }}')
    doc.add_paragraph('{% endfor %}')

    # Full Transcript
    doc.add_heading('Full Transcript', level=2)
    doc.add_paragraph('{% for seg in transcript %}[{{ seg.start }}] {{ seg.speaker }}: {{ seg.text }}\n{% endfor %}')

    path = os.path.join(output_dir, 'minutes_of_meeting.docx')
    doc.save(path)
    print(f"  [OK] {path}")


def create_one_on_one_template(output_dir):
    """1:1 Recap template."""
    doc = Document()
    doc.add_heading('1:1 Recap', level=1)
    doc.add_paragraph('{{ meeting.title }}')
    doc.add_paragraph('Date: {{ meeting.date }}  |  Duration: {{ meeting.duration }}')
    doc.add_paragraph('')
    doc.add_heading('Summary', level=2)
    doc.add_paragraph('{{ summary }}')
    doc.add_heading('Action Items', level=2)
    doc.add_paragraph('{% for item in action_items %}☐ [{{ item.owner }}] {{ item.description }}{% if item.due_date %} — Due: {{ item.due_date }}{% endif %}\n{% endfor %}')
    doc.add_heading('Discussion Notes', level=2)
    doc.add_paragraph('{% for seg in transcript %}{{ seg.speaker }}: {{ seg.text }}\n{% endfor %}')

    path = os.path.join(output_dir, 'one_on_one_recap.docx')
    doc.save(path)
    print(f"  [OK] {path}")


def create_standup_template(output_dir):
    """Standup Digest template."""
    doc = Document()
    doc.add_heading('Standup Digest', level=1)
    doc.add_paragraph('{{ meeting.date }}  |  Duration: {{ meeting.duration }}')
    doc.add_paragraph('')
    doc.add_heading('Summary', level=2)
    doc.add_paragraph('{{ summary }}')
    doc.add_heading('Action Items', level=2)
    doc.add_paragraph('{% for item in action_items %}• [{{ item.owner }}] {{ item.description }}\n{% endfor %}')
    doc.add_heading('Timeline', level=2)
    doc.add_paragraph('{% for event in timeline %}{{ event.timestamp }}s — {{ event.topic }}\n{% endfor %}')

    path = os.path.join(output_dir, 'standup_digest.docx')
    doc.save(path)
    print(f"  [OK] {path}")


def create_interview_template(output_dir):
    """Interview Notes template."""
    doc = Document()
    doc.add_heading('Interview Notes', level=1)
    doc.add_paragraph('{{ meeting.title }}')
    doc.add_paragraph('Date: {{ meeting.date }}  |  Duration: {{ meeting.duration }}')
    doc.add_paragraph('Participants: {% for person in attendees %}{{ person }}{% if not loop.last %}, {% endif %}{% endfor %}')
    doc.add_paragraph('')
    doc.add_heading('Summary', level=2)
    doc.add_paragraph('{{ summary }}')
    doc.add_heading('Key Points', level=2)
    doc.add_paragraph('{% for decision in decisions %}• {{ decision.description }} ({{ decision.speaker }})\n{% endfor %}')
    doc.add_heading('Full Transcript', level=2)
    doc.add_paragraph('{% for seg in transcript %}[{{ seg.start }}] {{ seg.speaker }}: {{ seg.text }}\n{% endfor %}')

    path = os.path.join(output_dir, 'interview_notes.docx')
    doc.save(path)
    print(f"  [OK] {path}")


def create_decision_log_template(output_dir):
    """Decision Log template."""
    doc = Document()
    doc.add_heading('Decision Log', level=1)
    doc.add_paragraph('{{ meeting.title }}  —  {{ meeting.date }}')
    doc.add_paragraph('')
    doc.add_heading('Decisions Made', level=2)
    doc.add_paragraph('{% for decision in decisions %}{{ loop.index }}. {{ decision.description }}\n   Decided by: {{ decision.speaker }}\n\n{% endfor %}')
    doc.add_heading('Related Action Items', level=2)
    doc.add_paragraph('{% for item in action_items %}• [{{ item.owner }}] {{ item.description }}\n{% endfor %}')

    path = os.path.join(output_dir, 'decision_log.docx')
    doc.save(path)
    print(f"  [OK] {path}")


def create_timeline_template(output_dir):
    """Timeline / Event Recap template."""
    doc = Document()
    doc.add_heading('Meeting Timeline', level=1)
    doc.add_paragraph('{{ meeting.title }}')
    doc.add_paragraph('Date: {{ meeting.date }}  |  Duration: {{ meeting.duration }}')
    doc.add_paragraph('')
    doc.add_heading('Summary', level=2)
    doc.add_paragraph('{{ summary }}')
    doc.add_heading('Event Timeline', level=2)
    doc.add_paragraph('{% for event in timeline %}⏱️ {{ event.timestamp }}s — [{{ event.phase }}] {{ event.topic }}\n{% endfor %}')
    doc.add_heading('Attendees', level=2)
    doc.add_paragraph('{% for person in attendees %}• {{ person }}\n{% endfor %}')

    path = os.path.join(output_dir, 'timeline_recap.docx')
    doc.save(path)
    print(f"  [OK] {path}")


if __name__ == "__main__":
    output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")
    os.makedirs(output_dir, exist_ok=True)

    print("Generating starter templates...")
    create_mom_template(output_dir)
    create_one_on_one_template(output_dir)
    create_standup_template(output_dir)
    create_interview_template(output_dir)
    create_decision_log_template(output_dir)
    create_timeline_template(output_dir)
    print(f"\nDone! {6} templates generated in: {output_dir}")
